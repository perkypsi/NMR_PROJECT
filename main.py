from utils import Signal, Spectrum, Data_spect, JShiftsHandler
import configparser
import numpy as np
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DPI = 70

# Инициализация конфиг файла
config = configparser.ConfigParser()
config.read('cfg.ini')

# Получаем значение параметра filename
filename = config.get('parameters', 'filename')

# Получаем значение параметра frequency
frequency = float(config.get('parameters', 'frequency'))

GROUPS = []

# Получаем группы атомов
for item in config.items('groups'):
    GROUPS.append(item[1].replace(' ', '').split(','))

# Обрабатываем out файл
data = JShiftsHandler(filename, groups=GROUPS)

# Обрабатываем каждую группу протонов
signals = []
for idx in range(len(data.groups)):

    # Константы спин-спинового взаимодействия
    j_constants = np.abs(np.array(data.j_groups[idx]))
    j_constants[idx] = 0

    # Количество атомов
    numbers_of_atoms = list(data.numbers_atoms.values())

    # Центр пика для данной группы
    shift = data.groups_shits[str(idx + 1)].iloc[0]

    # Подготовка мультиплетов
    base_intensity = numbers_of_atoms[idx] / np.max(numbers_of_atoms)
    signals.append(Signal(center=shift,
                          freq=frequency,
                          coupling=j_constants,
                          numbers_atoms=numbers_of_atoms,
                          base_intensity=base_intensity))


# Отрисовка всех сигналов
spectrum = Spectrum()
x = []
y = []
for signal in signals:
    for element in signal.signals:
        x.append(element[0])
        y.append(element[1])
# Подготовка полного спектра с добавлением нулей
for i in range(1000000):
    if i == 0:
        x.append(0)
    else:
        x.append(round(x[-1] + 20 / 1000000, 20))
    y.append(0)
# Преобразование данных
for idx, idy in zip(x, y):
    obj = Data_spect(idx, idy)
    spectrum.data.append(obj)

spectrum.sort_data()

x, y = spectrum.get_array()

# Отрисовка спектра
fig, axs = plt.subplots(figsize=(8, 6))
count = 0
axs.plot(x, y, linewidth=0.3, color='black')
axs.set_xlim(20, 0)
axs.grid(which='major', color='k', linestyle='-', alpha=0.2)
axs.grid(which='minor', linestyle='-', alpha=0.1)
axs.minorticks_on()
axs.grid(True)
# axs.invert_xaxis()
plt.xlabel("ppm", fontsize=14, fontweight='bold')
plt.ylabel("intensity", fontsize=14, fontweight='bold')
plt.show()


fig.savefig(f'{filename[:-4]}.png')

# Создаем отчет по файлу

document = Document()
style = document.styles['Normal']
style.font.name = 'Montserrat'
style.font.size = Pt(14)

# Добавление заголовка
document.add_heading('Отчет', 0)

# Добавление графика
document.add_picture(f'{filename[:-4]}.png', width=Inches(6))

p1 = document.add_paragraph().add_run('Рисунок 1. Спектр HNMR.')
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.italic = True
p1.font.size = Pt(14)

p2 = document.add_paragraph().add_run('Таблица 1.')
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p2.italic = True
p2.font.size = Pt(14)

p3 = document.add_paragraph().add_run('Химические сдвиги определенных групп атомов.')
p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
p3.italic = True
p3.font.size = Pt(14)

# Добавление химических сдвигов
shifts = document.add_table(data.groups_shits.shape[0]+1,
                            data.groups_shits.shape[1])

shifts.style = 'Table Grid'

for j in range(data.groups_shits.shape[-1]):
    shifts.cell(0,j).text = data.groups_shits.columns[j]

for i in range(data.groups_shits.shape[0]):
    for j in range(data.groups_shits.shape[-1]):
        shifts.cell(i+1,j).text = str(data.groups_shits.values[i,j])

for row in shifts.rows:
    for cell in row.cells:
        # Шрифт
        cell.paragraphs[0].runs[0].font.size = Pt(10)

# print(data.j_groups)

p4 = document.add_paragraph().add_run('Таблица 2.')
p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p4.italic = True
p4.font.size = Pt(14)

p5 = document.add_paragraph().add_run('Константы спин-спинового взаимодействия определенных групп атомов.')
p5.alignment = WD_ALIGN_PARAGRAPH.LEFT
p5.italic = True
p5.font.size = Pt(14)

# Добавление констант спин-спинового взаимодействия
constants = document.add_table(data.j_groups.shape[0]+1,
                            data.j_groups.shape[1]+1)

constants.style = 'Table Grid'
constants.alignment = WD_ALIGN_PARAGRAPH.CENTER
constants.style.font.size = Pt(10)

for j in range(data.j_groups.shape[-1]):
    constants.cell(0, j+1).text = str(data.j_groups.columns[j]+1)

for j in range(data.j_groups.shape[0]):
    constants.cell(j+1, 0).text = str(data.j_groups.columns[j]+1)

for i in range(data.j_groups.shape[0]):
    for j in range(data.j_groups.shape[-1]):
        constants.cell(i + 1,j + 1).text = str(data.j_groups.values[i,j])

for row in constants.rows:
    for cell in row.cells:
        # Шрифт
        # print(help(cell.paragraphs[0]))
        cell.paragraphs[0].style.font.size = Pt(8)


# Сохранение документа
document.save('report.docx')

